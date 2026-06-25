"""校历文本提取与本地规则分析。"""
from __future__ import annotations

import io
import json
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import urlparse

import requests

MAX_FILE_BYTES = 10 * 1024 * 1024
TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}


class CalendarImportError(RuntimeError):
    pass


def _limit(value: str) -> str:
    return value.replace("\x00", "").strip()[:60_000]


def _office(content: bytes, suffix: str) -> str:
    if suffix in {".xlsx", ".xls"}:
        from openpyxl import load_workbook
        book = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        rows = []
        for sheet in book.worksheets:
            rows.append("工作表：" + sheet.title)
            for row in sheet.iter_rows(values_only=True):
                cells = [str(x).strip() for x in row if x is not None and str(x).strip()]
                if cells:
                    rows.append(" | ".join(cells))
        return _limit("\n".join(rows))
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise CalendarImportError("缺少 python-docx，请执行 requirements 安装。") from exc
        doc = Document(io.BytesIO(content))
        rows = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            rows.extend(" | ".join(c.text.strip() for c in row.cells if c.text.strip()) for row in table.rows)
        return _limit("\n".join(rows))
    if suffix == ".pdf":
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        try:
            return _limit("\n".join(page.get_text() for page in doc))
        finally:
            doc.close()
    raise CalendarImportError("不支持的校历格式：{}".format(suffix or "未知格式"))


def _ocr(content: bytes, suffix: str) -> str:
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as exc:
        raise CalendarImportError("缺少 rapidocr_onnxruntime，请执行 requirements 安装。") from exc
    temp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
    try:
        temp.write(content)
        temp.close()
        result, _ = RapidOCR()(temp.name)
        return _limit("\n".join(item[1] for item in (result or []) if len(item) > 1))
    finally:
        Path(temp.name).unlink(missing_ok=True)


def extract_bytes(content: bytes, filename: str, content_type: str = "") -> Tuple[str, str]:
    if len(content) > MAX_FILE_BYTES:
        raise CalendarImportError("文件超过 10MB 限制。")
    suffix = Path(filename).suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        for encoding in ("utf-8-sig", "gb18030", "utf-8"):
            try:
                return _limit(content.decode(encoding)), "文本"
            except UnicodeDecodeError:
                continue
        raise CalendarImportError("文本编码无法识别。")
    if suffix in IMAGE_EXTENSIONS:
        return _ocr(content, suffix), "图片 OCR"
    if suffix in {".xlsx", ".xls", ".docx", ".pdf"}:
        try:
            return _office(content, suffix), "文件"
        except ImportError as exc:
            raise CalendarImportError("缺少文件解析依赖，请执行 requirements 安装。") from exc
        except Exception as exc:
            raise CalendarImportError("无法读取文件：{}".format(exc))
    if "html" in content_type.lower():
        try:
            from bs4 import BeautifulSoup
            return _limit(BeautifulSoup(content, "html.parser").get_text("\n", strip=True)), "网页"
        except ImportError as exc:
            raise CalendarImportError("缺少 beautifulsoup4，请执行 requirements 安装。") from exc
    raise CalendarImportError("请选择 TXT、Markdown、CSV、Excel、Word、PDF 或图片校历。")


def extract_file(path: str) -> Tuple[str, str]:
    source = Path(path)
    try:
        return extract_bytes(source.read_bytes(), source.name)
    except OSError as exc:
        raise CalendarImportError("无法读取文件：{}".format(exc))


def extract_url(url: str, timeout: int = 15) -> Tuple[str, str, str]:
    parsed = urlparse(url.strip())
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise CalendarImportError("只支持公开的 HTTP(S) 校历链接。")
    try:
        response = requests.get(url, timeout=timeout, stream=True, headers={"User-Agent": "CampusGrowthAssistant/0.1"})
        response.raise_for_status()
        chunks, size = [], 0
        for chunk in response.iter_content(8192):
            size += len(chunk)
            if size > MAX_FILE_BYTES:
                raise CalendarImportError("链接文件超过 10MB 限制。")
            chunks.append(chunk)
    except requests.RequestException as exc:
        raise CalendarImportError("无法下载校历链接：{}".format(exc))
    text, kind = extract_bytes(b"".join(chunks), Path(parsed.path).name or "calendar.html", response.headers.get("Content-Type", ""))
    return text, kind, url
